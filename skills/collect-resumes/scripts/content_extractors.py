# -*- coding: utf-8 -*-
"""
content_extractors.py — 统一内容提取层

支持格式：
  - PDF（PyMuPDF 提取文本）
  - 图片型 PDF（OCR，后端不可用时阻断）
  - DOCX（python-docx 提取正文/表格/页眉/页脚）
  - ZIP（安全展开后递归提取包内 PDF/DOCX）

统一返回 ExtractResult：
  text: 提取的文本
  nchar: 可提取字符数
  nimg: 图片数
  fmt: 检测到的格式
  warnings: 警告列表
  block_reason: 阻断原因（非 None = 不可放行）

"""

import os
import sys

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx as python_docx
except ImportError:
    python_docx = None


class ExtractResult:
    """统一提取结果。"""

    def __init__(self, text="", nchar=0, nimg=0, fmt="unknown",
                 warnings=None, block_reason=None):
        self.text = text
        self.nchar = nchar
        self.nimg = nimg
        self.fmt = fmt
        self.warnings = warnings or []
        self.block_reason = block_reason

    @property
    def blocked(self):
        return self.block_reason is not None

    def __repr__(self):
        return f"ExtractResult(fmt={self.fmt}, nchar={self.nchar}, nimg={self.nimg}, blocked={self.blocked})"


def _ocr_available():
    """检查 OCR 后端（tesseract 可执行文件）是否可用。"""
    import shutil
    return shutil.which("tesseract") is not None


def extract_pdf(path):
    """
    提取 PDF 文本。
    - 文本型 PDF：直接提取
    - 图片型 PDF（文本<20字且有图）：OCR；OCR 不可用 → 阻断
    - 加密/损坏 PDF → 阻断
    """
    if fitz is None:
        return ExtractResult(block_reason="PyMuPDF 不可用，无法提取 PDF")

    try:
        doc = fitz.open(path)
    except Exception as e:
        return ExtractResult(block_reason=f"PDF 打开失败（可能损坏/加密）: {e}")

    if doc.needs_pass:
        doc.close()
        return ExtractResult(block_reason="PDF 加密，无法提取")

    try:
        pages_text = []
        nimg = 0
        for pg in doc:
            pages_text.append(pg.get_text())
            nimg += len(pg.get_images())
        text = "\n".join(pages_text)
        nchar = len(text.strip())
        doc.close()

        if nchar < 20 and nimg > 0:
            # 图片型 PDF：需要 OCR
            if not _ocr_available():
                return ExtractResult(
                    text=text, nchar=nchar, nimg=nimg, fmt="pdf_scanned",
                    block_reason="图片型 PDF（无文本），tesseract OCR 后端不可用，无法验证姓名和薪酬"
                )
            # OCR 可用：尝试提取
            ocr_text = _ocr_pdf(path)
            if ocr_text is None:
                return ExtractResult(
                    text=text, nchar=nchar, nimg=nimg, fmt="pdf_scanned",
                    block_reason="OCR 提取失败或结果不足"
                )
            return ExtractResult(text=ocr_text, nchar=len(ocr_text.strip()), nimg=nimg, fmt="pdf_scanned")

        return ExtractResult(text=text, nchar=nchar, nimg=nimg, fmt="pdf")
    except Exception as e:
        return ExtractResult(block_reason=f"PDF 提取异常: {e}")


def _ocr_pdf(path):
    """对图片型 PDF 做 OCR。返回提取文本或 None（失败/不足）。"""
    try:
        import pytesseract
        from PIL import Image
        import io

        doc = fitz.open(path)
        texts = []
        for pg in doc:
            # 渲染页面为图片
            pix = pg.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            texts.append(pytesseract.image_to_string(img, lang="chi_sim+eng"))
        doc.close()
        result = "\n".join(texts)
        return result if len(result.strip()) >= 20 else None
    except Exception:
        return None


def extract_docx(path):
    """
    提取 DOCX 文本（正文、表格、页眉、页脚）。
    python-docx 不可用 → 阻断（不绕过）。
    """
    if python_docx is None:
        return ExtractResult(block_reason="python-docx 不可用，无法提取 DOCX")

    try:
        d = python_docx.Document(path)
    except Exception as e:
        return ExtractResult(block_reason=f"DOCX 打开失败（可能损坏/加密）: {e}")

    parts = []
    # 正文段落
    for para in d.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格
    for table in d.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    # 页眉页脚
    for section in d.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for para in section.footer.paragraphs:
            if para.text.strip():
                parts.append(para.text)

    text = "\n".join(parts)
    return ExtractResult(text=text, nchar=len(text.strip()), nimg=0, fmt="docx")


def extract(path):
    """
    根据文件扩展名调度到对应提取器。
    不支持的格式 → 阻断（不假放行）。
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
        return ExtractResult(block_reason="独立图片文件不直接验证，需先转为 PDF 或人工确认")
    return ExtractResult(block_reason=f"不支持的格式 .{ext}，无法提取验证")
